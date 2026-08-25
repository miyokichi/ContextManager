"""Word extractor: paragraphs grouped by heading into sections."""
from __future__ import annotations

from ..models import ExtractedDocument, ExtractedSection
from .base import ExtractionError


def extract(path, max_chars_per_section: int = 2000, max_sections: int = 40) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as e:
        raise ExtractionError("python-docx is required to read Word files") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ExtractionError(f"failed to open document: {e}") from e

    sections: list[ExtractedSection] = []
    current_heading = "Introduction"
    current_lines: list[str] = []

    def flush() -> None:
        text = "\n".join(current_lines).strip()
        if text:
            trimmed = text
            if len(trimmed) > max_chars_per_section:
                trimmed = trimmed[:max_chars_per_section] + "...(truncated)"
            sections.append(ExtractedSection(location=current_heading, text=trimmed))

    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        if style.startswith("Heading") or style == "Title":
            flush()
            current_heading = para.text.strip() or current_heading
            current_lines = []
            if len(sections) >= max_sections:
                break
        elif para.text.strip():
            current_lines.append(para.text)
    else:
        flush()

    if not sections:
        full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if len(full) > max_chars_per_section:
            full = full[:max_chars_per_section] + "...(truncated)"
        sections.append(ExtractedSection(location="Document", text=full))

    return ExtractedDocument(kind="docx", sections=sections)
